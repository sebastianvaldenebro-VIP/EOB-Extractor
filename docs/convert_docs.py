"""
Convert Markdown documentation to DOCX and PDF.

Usage:
    python3 convert_docs.py                          # converts TECHNICAL_HANDOFF.md
    python3 convert_docs.py ARCHITECTURE.md          # converts specific file
    python3 convert_docs.py --all                    # converts all *.md files in docs/

Asset for the /generate-project-docs command. Copied verbatim into <project>/docs/.
Do not edit per-project — edit this canonical copy at ~/.claude/skills/assets/convert_docs.py.
"""
import sys
import os
import glob
import re
import base64
import json

BASE = os.path.dirname(os.path.abspath(__file__))

_mermaid_cache = {}

def _mermaid_to_svg(diagram_code: str):
    try:
        import requests
        if diagram_code in _mermaid_cache:
            return _mermaid_cache[diagram_code]
        payload = json.dumps({"code": diagram_code, "mermaid": {"theme": "default"}})
        encoded = base64.urlsafe_b64encode(payload.encode()).decode().rstrip("=")
        url = f"https://mermaid.ink/svg/{encoded}"
        r = requests.get(url, timeout=30)
        if r.status_code == 200 and r.text.startswith("<svg"):
            svg = r.text
            svg = re.sub(r'style="max-width:[^"]*"', 'style="max-width:100%;height:auto"', svg)
            _mermaid_cache[diagram_code] = svg
            return svg
        print(f"  [mermaid.ink] HTTP {r.status_code}: {diagram_code[:40]!r}")
    except Exception as e:
        print(f"  [mermaid.ink] Error: {e}")
    return None

def _preprocess_markdown_for_pdf(md_content: str) -> str:
    def replace_mermaid(m):
        code = m.group(1).strip()
        svg = _mermaid_to_svg(code)
        if svg:
            print(f"  [mermaid] Rendered diagram ({len(code)} chars)")
            return f'<div class="diagram">{svg}</div>'
        escaped = code.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        return f'<pre class="mermaid-fallback"><code>{escaped}</code></pre>'
    return re.sub(r"```mermaid\s*\n(.*?)\n```", replace_mermaid, md_content, flags=re.DOTALL)

def convert_to_docx(md_file):
    from docx import Document
    from docx.shared import Pt, Inches
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    with open(md_file, "r") as f:
        lines = f.readlines()
    in_code_block = False
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        data_rows = [r for r in table_rows if not all(c.strip().replace("-","").replace("|","") == "" for c in r)]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        col_count = max(len(r) for r in data_rows)
        t = doc.add_table(rows=len(data_rows), cols=col_count)
        t.style = "Table Grid"
        for ri, row in enumerate(data_rows):
            for ci, cell_text in enumerate(row[:col_count]):
                cell = t.cell(ri, ci)
                cell.text = cell_text.strip()
                if ri == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph()
        in_table = False
        table_rows = []

    for line in lines:
        stripped = line.rstrip()
        if stripped.startswith("```"):
            if in_table:
                flush_table()
            in_code_block = not in_code_block
            if in_code_block:
                p = doc.add_paragraph()
                p.style = "No Spacing"
            continue
        if in_code_block:
            p = doc.add_paragraph(stripped)
            p.style = "No Spacing"
            run = p.runs[0] if p.runs else p.add_run(stripped)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            continue
        if stripped.startswith("|"):
            if not in_table:
                in_table = True
            cols = [c for c in stripped.split("|") if c != ""]
            table_rows.append(cols)
            continue
        else:
            if in_table:
                flush_table()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### ") and not stripped.startswith("#### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped == "---":
            doc.add_paragraph("─" * 60)
        elif stripped == "":
            doc.add_paragraph()
        else:
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            text = re.sub(r"\*(.*?)\*", r"\1", text)
            text = re.sub(r"`(.*?)`", r"\1", text)
            text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
            if text.startswith("- ") or text.startswith("* "):
                doc.add_paragraph(text[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", text):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", text), style="List Number")
            else:
                doc.add_paragraph(text)
    if in_table:
        flush_table()
    stem = os.path.splitext(os.path.basename(md_file))[0]
    out_path = os.path.join(BASE, f"{stem}.docx")
    doc.save(out_path)
    print(f"  DOCX saved: {out_path}")
    return out_path

def convert_to_pdf(md_file):
    try:
        import markdown
        from weasyprint import HTML
        from weasyprint.text.fonts import FontConfiguration
        with open(md_file, "r") as f:
            md_content = f.read()
        md_content = _preprocess_markdown_for_pdf(md_content)
        html_content = markdown.markdown(
            md_content,
            extensions=["tables", "fenced_code", "codehilite", "toc"]
        )
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
  @page {{ margin: 20mm 15mm; size: A4; }}
  body {{ font-family: Arial, sans-serif; font-size: 10pt; line-height: 1.5; color: #1a1a1a; word-wrap: break-word; overflow-wrap: break-word; }}
  h1 {{ font-size: 18pt; color: #1a365d; border-bottom: 2px solid #1a365d; padding-bottom: 6px; margin-top: 16px; }}
  h2 {{ font-size: 14pt; color: #2c5282; border-bottom: 1px solid #bee3f8; padding-bottom: 4px; margin-top: 20px; }}
  h3 {{ font-size: 11pt; color: #2b6cb0; margin-top: 14px; }}
  h4 {{ font-size: 10pt; color: #4a5568; margin-top: 10px; }}
  table {{ border-collapse: collapse; width: 100%; table-layout: fixed; margin: 10px 0; font-size: 8pt; page-break-inside: auto; }}
  th {{ background: #2c5282; color: white; padding: 5px 6px; text-align: left; word-wrap: break-word; overflow-wrap: break-word; white-space: normal; }}
  td {{ border: 1px solid #cbd5e0; padding: 4px 6px; vertical-align: top; word-wrap: break-word; overflow-wrap: break-word; white-space: normal; }}
  tr:nth-child(even) td {{ background: #f7fafc; }}
  code {{ background: #f0f4f8; padding: 1px 3px; border-radius: 2px; font-family: "Courier New", monospace; font-size: 8pt; word-break: break-all; }}
  pre {{ background: #1a202c; color: #e2e8f0; padding: 10px; border-radius: 4px; font-size: 7pt; white-space: pre-wrap; word-wrap: break-word; overflow-wrap: break-word; page-break-inside: avoid; }}
  pre code {{ background: none; color: inherit; padding: 0; word-break: normal; }}
  .diagram {{ margin: 12px 0; text-align: center; page-break-inside: avoid; }}
  .diagram svg {{ max-width: 100%; height: auto; }}
  .mermaid-fallback {{ background: #f0f4f8; color: #2d3748; border-left: 3px solid #4299e1; padding: 8px; font-size: 7pt; white-space: pre-wrap; word-wrap: break-word; }}
  blockquote {{ border-left: 3px solid #4299e1; margin: 0; padding: 4px 10px; color: #4a5568; }}
  a {{ color: #2b6cb0; word-break: break-all; }}
  hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 16px 0; }}
  ul, ol {{ padding-left: 18px; margin: 4px 0; }}
  li {{ margin: 2px 0; }}
  p {{ margin: 6px 0; }}
</style></head><body>{html_content}</body></html>"""
        stem = os.path.splitext(os.path.basename(md_file))[0]
        out_path = os.path.join(BASE, f"{stem}.pdf")
        font_config = FontConfiguration()
        HTML(string=full_html).write_pdf(out_path, font_config=font_config)
        print(f"  PDF  saved: {out_path}")
        return out_path
    except ImportError as e:
        print(f"  PDF skipped — missing dependency: {e}")
        return None

def convert_file(md_file):
    name = os.path.basename(md_file)
    print(f"\n--- Converting {name} ---")
    try:
        convert_to_docx(md_file)
    except Exception as e:
        print(f"  DOCX failed: {e}")
    try:
        convert_to_pdf(md_file)
    except Exception as e:
        print(f"  PDF  failed: {e}")

if __name__ == "__main__":
    args = sys.argv[1:]
    exclude = {"iam-remediation.md", "runbooks.md"}
    if "--all" in args:
        md_files = sorted(glob.glob(os.path.join(BASE, "*.md")))
        md_files = [f for f in md_files if os.path.basename(f).lower() not in exclude]
        for f in md_files:
            convert_file(f)
    elif args:
        target = args[0] if os.path.isabs(args[0]) else os.path.join(BASE, args[0])
        convert_file(target)
    else:
        convert_file(os.path.join(BASE, "TECHNICAL_HANDOFF.md"))
